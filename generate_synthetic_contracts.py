import os
import json
import glob
import subprocess
import cv2
import numpy as np
import re
from docx import Document
from PIL import Image
from pdf2image import convert_from_path

# --- Configuration ---
INPUT_DIR = "Employment Contracts" 
OUTPUT_DIR = "test data"  
DATASET_LABELS = "labels.json"

os.makedirs(OUTPUT_DIR, exist_ok=True)

def sanitize_docx(input_path: str, output_path: str):
    """
    Finds chains of Arabic (٠) or Western (0) zeros used as placeholders
    and replaces them with actual periods so LibreOffice renders them correctly.
    """
    doc = Document(input_path)
    # Match 3 or more consecutive zeros (Western or Eastern)
    zero_pattern = re.compile(r'[0٠]{3,}')
    
    # Clean standard paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = zero_pattern.sub('............', run.text)
                
    # Clean tables (contracts often use hidden tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = zero_pattern.sub('............', run.text)
                            
    doc.save(output_path)

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text for the ground_truth label."""
    doc = Document(file_path)
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            # Keep the Ground Truth string clean
            clean_text = re.sub(r'[\u200B-\u200F\u202A-\u202E]', '', text)
            lines.append(clean_text)
    return "\n".join(lines)

def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """Uses LibreOffice to silently convert a DOCX to a perfect PDF."""
    subprocess.run([
        "soffice", "--headless", "--convert-to", "pdf", 
        docx_path, "--outdir", output_dir
    ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
    
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    return os.path.join(output_dir, f"{base_name}.pdf")

def apply_degradations(img: Image.Image, quality: str) -> Image.Image:
    """Applies OpenCV transformations to simulate real-world scanning issues."""
    bgr = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
    
    if quality == "high":
        noise = np.random.normal(0, 2, bgr.shape).astype(np.uint8)
        bgr = cv2.add(bgr, noise)
    elif quality == "low_quality_blur":
        bgr = cv2.GaussianBlur(bgr, (11, 11), 3.0)
    elif quality == "low_quality_noise":
        noise = np.random.normal(0, 35, bgr.shape).astype(np.int16)
        bgr = np.clip(bgr.astype(np.int16) + noise, 0, 255).astype(np.uint8)
    elif quality == "low_quality_contrast":
        bgr = cv2.convertScaleAbs(bgr, alpha=0.5, beta=100)
    elif quality == "deskew_test":
        h, w = bgr.shape[:2]
        angle = np.random.uniform(-3.0, 3.0)
        M = cv2.getRotationMatrix2D((w/2, h/2), angle, 1)
        bgr = cv2.warpAffine(bgr, M, (w, h), borderValue=(255, 255, 255))

    return Image.fromarray(cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB))

def main():
    docx_files = glob.glob(os.path.join(INPUT_DIR, "*.docx"))
    if not docx_files:
        print(f"No DOCX files found in {INPUT_DIR}.")
        return

    dataset_labels = []

    for idx, doc_path in enumerate(docx_files):
        base_name = f"doc_{idx:03d}"
        print(f"Processing [{idx+1}/{len(docx_files)}]: {os.path.basename(doc_path)}")
        
        # 1. Create a sanitized temporary DOCX with proper dots
        sanitized_docx_path = os.path.join(OUTPUT_DIR, f"{base_name}_clean.docx")
        sanitize_docx(doc_path, sanitized_docx_path)
        
        # 2. Extract clean ground truth text from the sanitized version
        text = extract_text_from_docx(sanitized_docx_path)
        if not text:
            continue
            
        # 3. Convert the SANITIZED DOCX to PDF
        out_pdf = convert_docx_to_pdf(sanitized_docx_path, OUTPUT_DIR)
        
        # Rename the output PDF to our standardized naming scheme
        standard_pdf_path = os.path.join(OUTPUT_DIR, f"{base_name}_clean.pdf")
        if os.path.exists(out_pdf) and out_pdf != standard_pdf_path:
            os.rename(out_pdf, standard_pdf_path)
        
        dataset_labels.append({"filename": f"{base_name}_clean.pdf", "format": "pdf", "quality": "clean", "ground_truth": text})
        # Keep track of the original docx as well
        dataset_labels.append({"filename": f"{base_name}_clean.docx", "format": "docx", "quality": "clean", "ground_truth": text})

        # 4. Convert the clean PDF into a base image
        pages = convert_from_path(standard_pdf_path, dpi=300)
        if not pages:
            continue
        base_img = pages[0] 
        
        # 5. Generate degradations
        qualities = ["high", "low_quality_blur", "low_quality_noise", "low_quality_contrast", "deskew_test"]
        
        for q in qualities:
            degraded_img = apply_degradations(base_img.copy(), q)
            out_img_name = f"{base_name}_{q}.jpg"
            out_img_path = os.path.join(OUTPUT_DIR, out_img_name)
            degraded_img.save(out_img_path, "JPEG", quality=85)
            
            dataset_labels.append({
                "filename": out_img_name, 
                "format": "image", 
                "quality": q, 
                "ground_truth": text
            })

    # Save Labels
    labels_path = os.path.join(OUTPUT_DIR, DATASET_LABELS)
    with open(labels_path, 'w', encoding='utf-8') as f:
        json.dump(dataset_labels, f, ensure_ascii=False, indent=4)
        
    print(f"\n✅ Dataset generated successfully in '{OUTPUT_DIR}'!")

if __name__ == "__main__":
    main()